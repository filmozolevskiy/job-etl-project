"""Utility for extracting text from resume files (PDF and DOCX)."""

from __future__ import annotations

import logging
from io import BytesIO
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from shared.database import Database

    from .storage_service import StorageService

logger = logging.getLogger(__name__)

# Try to import PDF extraction libraries
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None  # type: ignore[assignment, misc]

# Try to import DOCX extraction library
try:
    from docx import Document
except ImportError:
    Document = None  # type: ignore[assignment, misc]


class ResumeTextExtractionError(Exception):
    """Raised when resume text extraction fails."""

    pass


def extract_text_from_resume(
    resume_id: int,
    user_id: int,
    storage_service: StorageService,
    database: Database,
    max_file_size: int = 10 * 1024 * 1024,  # 10MB default
) -> str:
    """Extract text content from a resume file.

    Args:
        resume_id: Resume ID to extract text from
        user_id: User ID (for ownership validation)
        storage_service: Storage service to read file
        database: Database connection to get resume metadata
        max_file_size: Maximum file size in bytes (default: 10MB)

    Returns:
        Extracted text content from the resume

    Raises:
        ResumeTextExtractionError: If extraction fails
        ValueError: If resume not found or user doesn't own it
        FileNotFoundError: If resume file doesn't exist
    """
    # Get resume metadata from database
    from .queries import GET_RESUME_BY_ID

    with database.get_cursor() as cur:
        cur.execute(GET_RESUME_BY_ID, (resume_id, user_id))
        result = cur.fetchone()
        if not result:
            raise ValueError(f"Resume {resume_id} not found or access denied")

        if cur.description is None:
            raise ValueError("No description available from cursor")
        columns = [desc[0] for desc in cur.description]
        resume_data = dict(zip(columns, result))

    file_path = resume_data.get("file_path")
    if not file_path:
        raise ResumeTextExtractionError("Resume has no file path")

    # Validate file path to prevent path traversal attacks
    if ".." in file_path or file_path.startswith("/"):
        raise ResumeTextExtractionError(f"Invalid file path: {file_path}")

    # Validate file size before reading
    file_size = resume_data.get("file_size", 0)
    if file_size > max_file_size:
        raise ResumeTextExtractionError(
            f"Resume file too large: {file_size} bytes (max: {max_file_size} bytes)"
        )

    # Determine file type from path
    file_ext = file_path.lower().rsplit(".", 1)[-1] if "." in file_path else ""

    # Read file content
    try:
        file_content = storage_service.get_file(file_path)

        # Double-check file size after reading (in case database value is wrong)
        if len(file_content) > max_file_size:
            raise ResumeTextExtractionError(
                f"Resume file too large: {len(file_content)} bytes (max: {max_file_size} bytes)"
            )
    except FileNotFoundError as e:
        raise FileNotFoundError(f"Resume file not found: {file_path}") from e
    except Exception as e:
        if isinstance(e, ResumeTextExtractionError):
            raise
        raise ResumeTextExtractionError(f"Failed to read resume file: {e}") from e

    # Extract text based on file type
    if file_ext == "pdf":
        return _extract_pdf_text(file_content)
    elif file_ext == "docx":
        return _extract_docx_text(file_content)
    else:
        raise ResumeTextExtractionError(
            f"Unsupported file type: {file_ext}. Supported types: pdf, docx"
        )


def _extract_pdf_text(file_content: bytes) -> str:
    """Extract text from PDF file content.

    Args:
        file_content: Binary PDF file content

    Returns:
        Extracted text from PDF

    Raises:
        ResumeTextExtractionError: If PDF extraction fails
    """
    if PyPDF2 is None:
        raise ResumeTextExtractionError(
            "PyPDF2 library is not installed. Install with: pip install PyPDF2>=3.0.0"
        )

    try:
        pdf_file = BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text_parts = []

        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            except Exception as e:
                logger.warning(f"Failed to extract text from PDF page {page_num + 1}: {e}")
                continue

        if not text_parts:
            raise ResumeTextExtractionError("No text could be extracted from PDF")

        return "\n\n".join(text_parts)

    except Exception as e:
        if isinstance(e, ResumeTextExtractionError):
            raise
        raise ResumeTextExtractionError(f"Failed to extract text from PDF: {e}") from e


def _extract_docx_text(file_content: bytes) -> str:
    """Extract text from DOCX file content.

    Args:
        file_content: Binary DOCX file content

    Returns:
        Extracted text from DOCX

    Raises:
        ResumeTextExtractionError: If DOCX extraction fails
    """
    if Document is None:
        raise ResumeTextExtractionError(
            "python-docx library is not installed. Install with: pip install python-docx>=1.1.0"
        )

    try:
        docx_file = BytesIO(file_content)
        doc = Document(docx_file)
        text_parts = []

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_texts.append(cell.text.strip())
                if row_texts:
                    text_parts.append(" | ".join(row_texts))

        if not text_parts:
            raise ResumeTextExtractionError("No text could be extracted from DOCX")

        return "\n\n".join(text_parts)

    except Exception as e:
        if isinstance(e, ResumeTextExtractionError):
            raise
        raise ResumeTextExtractionError(f"Failed to extract text from DOCX: {e}") from e
